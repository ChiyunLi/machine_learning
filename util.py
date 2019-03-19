#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sun Mar 11 15:33:20 2018

@author: slave1
"""

from sklearn.metrics import confusion_matrix, auc, accuracy_score
import pandas as pd
import numpy as np
import statsmodels.api as sm

# ----設定繪圖-------
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties 
import seaborn as sns 

######################## util #####################################

def docx_table(df,doc):    
    # add a table to the end and create a reference variable
#  extra row is so we can add the header row
    doc.add_heading('測試文件', 1)
    t = doc.add_table(df.shape[0]+1, df.shape[1])
    
#    # add the header rows.
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
        
#    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])


def logistic_model(X_train,y_train,X_test,y_test,
                    sales_price = 3500,
                    marketing_expense = 300,
                    product_cost = 1800,
                    plot_name = 'logistic_regression'
                   ):
    
    X_train_log = X_train.copy()
    X_train_log['intercept'] = 1
    logistic = sm.Logit(y_train,X_train_log)
    
    # fit the model
    result = logistic.fit()
    result_df = results_summary_to_dataframe(result, plot_name = plot_name)
    
    
    X_test_log = X_test.copy()
    X_test_log['intercept'] = 1
    y_test_df=pd.DataFrame(y_test)
    y_test_df[plot_name+'_pred'] = result.predict(X_test_log)
    y_test_df['pred_yn']= np.where(y_test_df[plot_name+'_pred']>=0.5, 1,0)
    
    
    conf_logist = confusion_matrix(y_test_df['buy'], y_test_df['pred_yn'])
    
    plot_confusion_matrix(conf_logist, ['no','buy'],
                          title=plot_name+"Confusion Matrix plot", cmap=plt.cm.Reds)#, cmap=plt.cm.Reds
    
    # model_profit 
    model_profit = sales_price * conf_logist[1,1] - conf_logist[::,1].sum() * marketing_expense - product_cost * conf_logist[1,1]
    
    model_profit_df = pd.DataFrame({
            '項目' : ['單品價格', '單品營業成本', '單品行銷費用', '利潤'],
            '金額' : [sales_price,product_cost, marketing_expense, '-'],
            '目標對象' : [conf_logist[1,1],conf_logist[1,1], conf_logist[::,1].sum(), '-'],
            '小計' : [sales_price* conf_logist[1,1], product_cost* conf_logist[1,1], marketing_expense * conf_logist[::,1].sum(),model_profit  ],
            })
    
    
    # all_profit 
    all_profit = sales_price*conf_logist[1,::].sum() - product_cost* conf_logist[1,::].sum()- marketing_expense *  conf_logist.sum()
    
    all_df = pd.DataFrame({
            '項目' : ['單品價格', '單品營業成本', '單品行銷費用', '利潤'],
            '金額' : [sales_price,product_cost, marketing_expense, '-'],
            '目標對象' : [conf_logist[1,::].sum(), conf_logist[1,::].sum(), conf_logist.sum(), '-'],
            '小計' : [sales_price*conf_logist[1,::].sum(), product_cost* conf_logist[1,::].sum(), marketing_expense *  conf_logist.sum(),all_profit  ],
            })
    
    
    # -------single model summary--------
    
    print( "################ summary ################ ")
    
    print(confusion_matrix(y_test_df['buy'], y_test_df['pred_yn']))
#    print("____________________{}分類報告____________________".format(plot_name))
#    print(classification_report(y_test_df['buy'], y_test_df['pred_yn']))
    print(accuracy_score(y_test_df['buy'], y_test_df['pred_yn']))      
    
          
    # importance
    print( '------------------ 應注意之變數 ------------------' )
    print('、'.join(result_df['變數'].tolist()))
    print('\n'.join(result_df['意涵'].tolist()))
    
    
    # profit comparison
    if model_profit - all_profit > 0 :
        print( '------------------模型相對權市場行銷來說【賺錢】------------------' )
        print( '模型比全市場行銷賺 $' + str(model_profit - all_profit) )
        print( '比較全市場行銷來說，淨利減少' + str( round(model_profit / all_profit, 3) ) + '倍')
        
    else:
        print( '------------------模型相對權市場行銷來說【賠錢】------------------' )
        print( '模型比全市場行銷損失 $' + str(model_profit - all_profit) )
        print( '比較全市場行銷來說，淨利增加' + str( round(model_profit / all_profit, 3) ) + '倍')
    
    print( '------------------全市場行銷利潤矩陣------------------' )
    print(all_df)
    all_df
    all_df.to_excel(plot_name+'全市場行銷利潤矩陣.xlsx')
    print('全市場行銷利潤矩陣.xlsx saved')
    
    print( '------------------' +plot_name+ '模型行銷利潤矩陣------------------' )
    print(model_profit_df)
    model_profit_df.to_excel(plot_name+'模型行銷利潤矩陣.xlsx')
    print('模型行銷利潤矩陣.xlsx saved')
    
    
    print( '------------------' +plot_name+ '重要變數表------------------' )
#    print(result_df)
    result_df.to_excel(plot_name+'重要變數表.xlsx')
    return all_df, model_profit_df, result_df,y_test_df
    

def logistic_importance(
        X_train,
        y_train,
        X_test,
        y_test,
        plot_name
                   ):
        
    X_train_log = X_train.copy()
    X_train_log['intercept'] = 1
    logistic = sm.Logit(y_train,X_train_log)
    
    # fit the model
    result = logistic.fit()
    result_df = results_summary_to_dataframe(result, plot_name = plot_name)
    return result_df 


def logistic_conf(
        X_train,
        y_train,
        X_test,
        y_test,
        plot_name
                   ):
        
    X_train_log = X_train.copy()
    X_train_log['intercept'] = 1
    logistic = sm.Logit(y_train,X_train_log)
    
    # fit the model
    result = logistic.fit()
#    result_df = results_summary_to_dataframe(result, plot_name = plot_name)
        
        
        
    X_test_log = X_test.copy()
    X_test_log['intercept'] = 1
    y_test_df=pd.DataFrame(y_test)
    y_test_df['pred'] = result.predict(X_test_log)
    y_test_df['pred_yn']= np.where(y_test_df['pred']>=0.5, 1,0)
    
    
    conf_logist = confusion_matrix(y_test_df['buy'], y_test_df['pred_yn'])
    
    plot_confusion_matrix(conf_logist, ['no','buy'],
                          title=plot_name+"Confusion Matrix plot", cmap=plt.cm.Reds)#, cmap=plt.cm.Reds
    
    
    print( "################ summary ################ ")
    print('Confusion matrix')
    print(confusion_matrix(y_test_df['buy'], y_test_df['pred_yn']))
#    print("____________________{}分類報告____________________".format(plot_name))
#    print(classification_report(y_test_df['buy'], y_test_df['pred_yn']))
    print("Test Accuracy = {:.3f}".format(accuracy_score(y_test_df['buy'], y_test_df['pred_yn'])))
    
    
    return conf_logist 




def detect_str_columns(data):
    '''
    1. 偵測有字串的欄位
    2. 挑選出來，準備encoding
    
    '''
    strlist = list(set(np.where((data.applymap(type)==str))[1].tolist()))
    return data.columns[strlist].tolist()


def results_summary_to_dataframe(results,plot_name):
    '''This takes the result of an statsmodel results table and transforms it into a dataframe'''
#    print(results.summary())
    pvals = round(results.pvalues, 3)
    coeff = round(results.params, 6)
#    conf_lower = round(results.conf_int()[0], 3)
#    conf_higher = round(results.conf_int()[1], 3)

    results_df = pd.DataFrame({
                               "參數":coeff,
                               "p_values":pvals,
#                               "conf_lower":conf_lower,
#                               "conf_higher":conf_higher
                                })
    
    
    results_df = results_df[results_df['p_values']<0.05]
    results_df = results_df.reset_index()
    results_df.columns = ['變數', '參數', 'p_values']
    #Reordering...
#    results_df = results_df[["coeff","pvals","conf_lower","conf_higher"]]
    
    feat_imp = results_df.參數
    feat = results_df['變數'].tolist()
    res_data = pd.DataFrame({'Features': feat, 'Importance': feat_imp}).sort_values(by='Importance', ascending=True)
#    res_data.plot('Features', 'Importance', kind='bar', title='Feature Importances')
    plt.figure(figsize=(10,10))
    plt.bar(res_data['Features'], res_data['Importance'])
    plt.title('Importance')
#    plt.subplots_adjust(left=0.7, right=0.8, top=0.9, bottom=0.7)
    plt.ylabel(plot_name+' Feature Importance Score')
    plt.savefig(plot_name+'.png', dpi=300)
    plt.show()
    
    results_df['轉換後參數'] = round(np.exp(results_df['參數']), 3)
    results_df = results_df.sort_values(['轉換後參數'],ascending = False)
    results_df['程度排名'] =range(1,len(results_df)+1 ) 
    results_df['意涵'] = '每增加 「' + results_df['變數'] +  '」 的1個單位 ，等同增加' + results_df['轉換後參數'].astype(str) + '的可能購買倍數'
    
    return results_df

import itertools

def plot_confusion_matrix(cm, classes,
                          normalize=False,
                          title='Confusion matrix',
                          cmap=plt.cm.Blues):
    """
    This function prints and plots the confusion matrix.
    Normalization can be applied by setting `normalize=True`.
    """
    fig = plt.figure()
    plt.imshow(cm, interpolation='nearest', cmap=cmap)
    plt.title(title)
    plt.colorbar()
    tick_marks = np.arange(len(classes))
    plt.xticks(tick_marks, classes, rotation=0)
    plt.yticks(tick_marks, classes)

    if normalize:
        cm = cm.astype('float') / cm.sum(axis=1)[:, np.newaxis]
        #print("Normalized confusion matrix")
    else:
        1#print('Confusion matrix, without normalization')

    #print(cm)

    thresh = cm.max() / 2.
    for i, j in itertools.product(range(cm.shape[0]), range(cm.shape[1])):
        plt.text(j, i, cm[i, j],
                 horizontalalignment="center",
                 color="white" if cm[i, j] > thresh else "black")

    plt.tight_layout()
    plt.ylabel('True label')
    plt.xlabel('Predicted label')
    fig.savefig(title+'.png', dpi=300)
# make demographics variables dummy 

def get_dummies(dummy, dataset):
    ''''
    make variables dummies
    ref：http://blog.csdn.net/weiwei9363/article/details/78255210
    '''
    dummy_fields = list(dummy)
    for each in dummy_fields:
        dummies = pd.get_dummies( dataset.loc[:, each], prefix=each ) 
        dataset = pd.concat( [dataset, dummies], axis = 1 )
    
    fields_to_drop = dummy_fields
    dataset = dataset.drop( fields_to_drop, axis = 1 )
    return dataset

# transform_fb_date
def transform_fb_date(data = None, columns = 'created_time', splitby = '-'):
    
    data[columns] = data[columns].str.replace('T','-')
    data[columns] = data[columns].str.replace(':','-')
    data[columns] = data[columns].str.split('+').str[0]
    
    col_num =  data[columns][0].count('-')+1
    date = data[columns].str.split('-', col_num, expand=True)
    date.columns = ['year', 'month', 'day', 'hour', 'min', 'sec']
    
    combine= pd.concat([data, date ], axis = 1) #keys=[],names=['gg', 'example'],
    #combine.columns = [columns + '_cut', columns]
    
    # drop original col
    fields_to_drop = [columns]
    combine = combine.drop( fields_to_drop, axis = 1 )
    return combine


#
from docx import Document
from docx.shared import Cm


def model_testRF(clf, X_train,y_train,X_test,y_test,sales_price = 3500,
                    marketing_expense = 300,
                    product_cost = 1800,
                    plot_name = 'logistic_regression') :
    model = clf.fit(X_train, y_train)
    y_pred = model.predict(X_test)
    y_pred_prob = model.predict_proba(X_test)[:,1]
    
    y_test_df=pd.DataFrame(y_test)
    y_test_df[plot_name+'_pred'] = y_pred_prob
    
    #Confusion Matrix
    conf_logist = confusion_matrix(y_test, y_pred)
    
    # 畫conf matrix
    plot_confusion_matrix(conf_logist, ['no','buy'],
                          title=plot_name+"Confusion Matrix plot", cmap=plt.cm.Reds)#, cmap=plt.cm.Reds
    
    # -------single model summary--------
    

    print( "################ summary ################ ")
    
    print(confusion_matrix(y_test, y_pred))
#    print("____________________{}分類報告____________________".format(plot_name))
#    print(classification_report(y_test, y_pred))
    print("Training Accuracy = {:.3f}".format(model.score(X_train, y_train)))
    print("Test Accuracy = {:.3f}".format(model.score(X_test, y_test)))

def model(clf, X_train,y_train,X_test,y_test,sales_price = 3500,
                    marketing_expense = 300,
                    product_cost = 1800,
                    plot_name = 'logistic_regression') :
    model = clf.fit(X_train, y_train)
    y_pred = model.predict(X_test)
    y_pred_prob = model.predict_proba(X_test)[:,1]
    
    y_test_df=pd.DataFrame(y_test)
    y_test_df[plot_name+'_pred'] = y_pred_prob
    
    #Confusion Matrix
    conf_logist = confusion_matrix(y_test, y_pred)
    
    # 畫conf matrix
    plot_confusion_matrix(conf_logist, ['no','buy'],
                          title=plot_name+"Confusion Matrix plot", cmap=plt.cm.Reds)#, cmap=plt.cm.Reds
    
    feat_imp = model.feature_importances_
    feat = X_train.columns.tolist()
    res_data = pd.DataFrame({'Features': feat, 'Importance': feat_imp}).sort_values(by='Importance', ascending=True)
    res_data.plot('Features', 'Importance', kind='barh', title='Feature Importances',stacked=True, figsize = (15,10))
    plt.ylabel(plot_name+' Feature Importance Score')
    plt.savefig(plot_name+'.png', dpi=300)
    plt.show()
    
    
    res_data = res_data.sort_values('Importance',ascending=False)
    
    # > 50%
    res_data = res_data[res_data['Importance']>res_data['Importance'].describe()['50%']]
    
    
    
    # model_profit 
    model_profit = sales_price * conf_logist[1,1] - conf_logist[::,1].sum() * marketing_expense - product_cost * conf_logist[1,1]
    
    model_profit_df = pd.DataFrame({
            '項目' : ['單品價格', '單品營業成本', '單品行銷費用', '利潤'],
            '金額' : [sales_price,product_cost, marketing_expense, '-'],
            '目標對象' : [conf_logist[1,1],conf_logist[1,1], conf_logist[::,1].sum(), '-'],
            '小計' : [sales_price* conf_logist[1,1], product_cost* conf_logist[1,1], marketing_expense * conf_logist[::,1].sum(),model_profit  ],
            })
    
    
    # all_profit 
    all_profit = sales_price*conf_logist[1,::].sum() - product_cost* conf_logist[1,::].sum()- marketing_expense *  conf_logist.sum()
    
    all_df = pd.DataFrame({
            '項目' : ['單品價格', '單品營業成本', '單品行銷費用', '利潤'],
            '金額' : [sales_price,product_cost, marketing_expense, '-'],
            '目標對象' : [conf_logist[1,::].sum(), conf_logist[1,::].sum(), conf_logist.sum(), '-'],
            '小計' : [sales_price*conf_logist[1,::].sum(), product_cost* conf_logist[1,::].sum(), marketing_expense *  conf_logist.sum(),all_profit  ],
            })

    # -------single model summary--------
    
    # 創造docx
    doc = Document()
    
    # save the doc

    print( "################ summary ################ ")
    doc.add_heading("################ summary ################ ", 0)
    
    print(confusion_matrix(y_test, y_pred))
    print("Training Accuracy = {:.3f}".format(model.score(X_train, y_train)))
    print("Test Accuracy = {:.3f}".format(model.score(X_test, y_test)))
    doc.add_paragraph("Training Accuracy = {:.3f}".format(model.score(X_train, y_train)))
    doc.add_paragraph("Test Accuracy = {:.3f}".format(model.score(X_test, y_test)))    
    
    
    # importance
    print( '------------------ 應注意之變數 ------------------' )
    doc.add_heading('------------------ 應注意之變數 ------------------', 1)
                    
    print(res_data)
    print('\n'.join(res_data["Features"].tolist()))
    docx_table(res_data,doc )
    doc.add_paragraph('\n'.join(res_data["Features"].tolist()))
    
    # profit comparison
    if model_profit - all_profit > 0 :
        print( '------------------模型相對權市場行銷來說【賺錢】------------------' )
        print( '模型比全市場行銷賺 $' + str(model_profit - all_profit) )
        
        doc.add_heading('------------------模型相對權市場行銷來說【賺錢】------------------',1)
        doc.add_paragraph( '模型比全市場行銷賺 $' + str(model_profit - all_profit) )
        
        if all_profit<0:
            all_profit2 = model_profit - all_profit
            print( '比較全市場行銷來說，行銷費用少' + str( round(model_profit / all_profit2, 3) ) + '倍')
        else:
            print( '比較全市場行銷來說，淨利增加' + str( round(model_profit / all_profit, 3) ) + '倍')
        
    else:
        print( '------------------模型相對權市場行銷來說【賠錢】------------------' )
        print( '模型比全市場行銷損失 $' + str(model_profit - all_profit) )
        print( '比較全市場行銷來說，淨利減少' + str( round(model_profit / all_profit, 3) ) + '倍')
        doc.add_heading('------------------模型相對權市場行銷來說【賺錢】------------------',1)
        doc.add_paragraph( '模型比全市場行銷賺 $' + str(model_profit - all_profit) )
        
        
        
        
    print( '------------------全市場行銷利潤矩陣------------------' )
    print(all_df)
    all_df.to_excel(plot_name+'全市場行銷利潤矩陣.xlsx')
    print('全市場行銷利潤矩陣.xlsx saved')
    
    doc.add_heading('------------------全市場行銷利潤矩陣------------------',1)
    docx_table(all_df,doc )
        
    
    
    
    print( '------------------' +plot_name+ '模型行銷利潤矩陣------------------' )
    print(model_profit_df)
    model_profit_df.to_excel(plot_name+'模型行銷利潤矩陣.xlsx')
    print('模型行銷利潤矩陣.xlsx saved')
    
    doc.add_heading( '------------------' +plot_name+ '模型行銷利潤矩陣------------------',1)
    docx_table(model_profit_df,doc )
        
    
    
    print( '------------------' +plot_name+ '重要變數表------------------' )
    print(result_df)
    res_data.to_excel(plot_name+'重要變數表.xlsx')
    
    
    doc.add_heading(  '------------------' +plot_name+ '重要變數表------------------',1)
    docx_table(res_data,doc )
    
    # pic
    doc.add_picture(plot_name+'.png')
    doc.add_picture(plot_name+"Confusion Matrix plot.png")
#
    doc.save(plot_name+'.docx')
#    
    return all_df, model_profit_df, res_data, y_test_df
    


# part 獲利閥值長條圖
def profit_linechart(y_test_df,
                     
                    sales_price = 3500,
                    marketing_expense = 300,
                    product_cost = 1800,
                    plot_name = 'logistic_regression' ):
    
    
    profit_line = []
    for i in np.arange(0,1,0.01):
        
        # set threshold
        y_test_df['pred_yn']= np.where(y_test_df[y_test_df.columns[1]]>=i, 1,0)    
        
        
        conf_logist = confusion_matrix(y_test_df['buy'], y_test_df['pred_yn'])
        
        # model_profit 
        model_profit = sales_price * conf_logist[1,1] - conf_logist[::,1].sum() * marketing_expense - product_cost * conf_logist[1,1]
        
        
        # all_profit 
        all_profit = sales_price*conf_logist[1,::].sum() - product_cost* conf_logist[1,::].sum()- marketing_expense *  conf_logist.sum()
        
        # 將所有threshold append在一起
        profit_line.append([ i,model_profit, all_profit])
    
    profit_line = pd.DataFrame(profit_line, columns= ['閥值', plot_name, '全市場'])
    
    # draw
    X_max = profit_line[profit_line[plot_name] ==profit_line[plot_name].max()]['閥值']
    Y_max =  profit_line[plot_name].max()
    
    profit_line = profit_line.rename( columns= {'閥值': 'threshold', '全市場' : 'all_market'})
    profit_line.plot.line(x='threshold', y=[plot_name,     'all_market'],figsize=(15,10))
    plt.scatter(X_max,Y_max, c='red', marker='o',alpha=0.5)
    plt.text(X_max-0.005, Y_max+10000, plot_name+' best profit$ ' + str(Y_max ) + ', threshold='+  str(X_max.values[0]))
    plt.ylabel('expected profit')
    plt.savefig(plot_name+'_預期獲利最佳化模型與閥值折線圖.png', dpi=300)
    plt.show()
    
    document = Document()
    document.add_heading(plot_name+'_預期獲利最佳化模型與閥值折線圖.png', 0)
    document.add_paragraph('台灣第一個行銷資料科學(MDS)知識部落\n\n本粉絲專頁在探討資料科學之基礎概念、趨勢、新工具和實作，讓粉絲們瞭解資料科學的行銷運用,並開啟厚植數據分析能力之契機')
    document.add_picture(plot_name+'_預期獲利最佳化模型與閥值折線圖.png')
    document.save(plot_name+'_預期獲利最佳化模型與閥值折線圖')




def cut_off_calu(y_test_df,
                    sales_price = 3500,
                    marketing_expense = 300,
                    product_cost = 1800,
                    plot_name = 'logistic_regression'):
    profit_line = []
    for i in np.arange(0,1,0.01):
        
        # set threshold
        y_test_df['pred_yn']= np.where(y_test_df[plot_name]>=i, 1,0)    
        
        
        conf_logist = confusion_matrix(y_test_df['buy'], y_test_df['pred_yn'])
        
        # model_profit 
        model_profit = sales_price * conf_logist[1,1] - conf_logist[::,1].sum() * marketing_expense - product_cost * conf_logist[1,1]
        
        
        # all_profit 
        all_profit = sales_price*conf_logist[1,::].sum() - product_cost* conf_logist[1,::].sum()- marketing_expense *  conf_logist.sum()
        
        # 將所有threshold append在一起
        profit_line.append([ i,model_profit, all_profit])
    
    profit_line = pd.DataFrame(profit_line, columns= ['閥值', plot_name, '全市場'])
    return profit_line 
    
# part 獲利閥值長條圖
def profit_linechart_all(y_test_df ,
                    sales_price = 3500,
                    marketing_expense = 300,
                    product_cost = 1800
                    ):
    
    allcon  = []
    for i in y_test_df:
            
        profit_line = cut_off_calu(i,
                        sales_price = sales_price,
                        marketing_expense = marketing_expense,
                        product_cost = product_cost,
                        plot_name = i.columns[1])
        allcon.append(profit_line)
        
    
    from functools import reduce
    allcondf= reduce(lambda x, y: pd.merge(x, y, on = ['閥值', '全市場']), allcon)
        
    # draw
    
    
        
    allcondf = allcondf.rename( columns= {'閥值': 'threshold', '全市場' : 'all_market'})
#    profit_line.plot.line(x='threshold', y=[plot_name,     'all_market'],figsize=(15,10))
    
    allcondf.plot.line(x='threshold', y=allcondf.drop(columns='threshold').columns.tolist(), figsize=(15,10))
    
    for plot_name in allcondf.drop(columns='threshold').columns.tolist():
        
        # 因為全市場沒有最高點
        if not plot_name=='all_market':
            X_max = allcondf[allcondf[plot_name] ==allcondf[plot_name].max()]['threshold']
            Y_max =  allcondf[plot_name].max()
            plt.scatter(X_max,Y_max, c='red', marker='o',alpha=0.5)
            
            if 'Random' in plot_name:
                
                plt.text(X_max-0.005, Y_max+Y_max * 0.07 , plot_name.replace('_pred','')+' best profit $ ' + str(Y_max ) + ', threshold='+  str(X_max.values[0]))
    
            elif 'xgb' in plot_name:
                plt.text(X_max-0.005, Y_max+Y_max * 0.02, plot_name.replace('_pred','')+' best profit $ ' + str(Y_max ) + ', threshold='+  str(X_max.values[0]))
            
            else:
                plt.text(X_max-0.005, Y_max + Y_max * 0.02 , plot_name.replace('_pred','')+' best profit $ ' + str(Y_max ) + ', threshold='+  str(X_max.values[0]))
    
            
    
    plt.ylabel('expected profit')
    plt.savefig('預期獲利最佳化模型與閥值折線圖'+'.png', dpi=300)
    plt.show()

